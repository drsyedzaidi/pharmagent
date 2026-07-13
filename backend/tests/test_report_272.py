"""ICH M4E Module 2.7.2 report generation tests.

Covers:
  - Empty state  → produces DOCX with placeholders, no crash
  - NCA-only     → Table 2.7.2.1 present; popPK section is placeholder
  - NLME present → Table 2.7.2.2 populated; conclusions mention model
  - SCM present  → Table 2.7.2.3 populated; special populations text
  - Full state   → all sections populated; file written to disk
  - StudyInfo    → args propagate into cover + background section
  - HTTP endpoint → POST /api/sessions/{sid}/report/272 returns report_path
  - PharmState   → study_info + regulatory_report_path written correctly
"""
from __future__ import annotations

import itertools
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.core.orchestrator import Orchestrator
from app.core.pharmstate import PharmState, StudyInfo
from app.core.llm import MockLLM
from app.core.store import SessionStore
from app.tools.base import ToolContext
from app.tools.regulatory_tools import generate_272


SAMPLE = str(Path(__file__).parent.parent / "sample_data" / "oral_pk.csv")


# ── helpers ───────────────────────────────────────────────────────────────────

def _ctx(tmp_path: Path) -> ToolContext:
    return ToolContext(data_dir=str(tmp_path), dataset_store=None)  # type: ignore[arg-type]


def _orch():
    counter = itertools.count()
    return Orchestrator(llm=MockLLM(), clock=lambda: f"t{next(counter)}",
                        store=SessionStore(":memory:"))


def _base_state() -> PharmState:
    return PharmState(
        dataset_metadata={"dataset_id": "test_ds", "n_subjects": 12,
                          "dose_levels": "100 mg", "dataset_sha256": "abc123"},
    )


def _nca_state() -> PharmState:
    st = _base_state()
    return st.model_copy(update={
        "nca_parameters": [
            {"subject": 1, "dose": 100, "Cmax": 450.0, "AUC_inf": 3200.0,
             "t_half": 8.5, "CL_F": 2.7, "Vz_F": 33.0}
        ],
        "nca_summary": {
            "by_dose": [{"dose": 100, "n": 12,
                         "Cmax_geomean": 440.0, "Cmax_geocv_pct": 28.0,
                         "AUC_inf_geomean": 3150.0, "AUC_inf_geocv_pct": 32.0,
                         "t_half_median": 8.4, "CL_F_geomean": 2.75}]
        },
    })


def _nlme_state(base: PharmState | None = None) -> PharmState:
    st = base or _nca_state()
    return st.model_copy(update={
        "nlme_results": {
            "status": "ok", "method": "focei",
            "label": "oral_1cmt", "iiv_params": ["CL", "V"],
            "error_model": "proportional", "ofv": -234.5,
            "converged": True, "condition_number": 38.2,
            "theta": {"CL": 2.71, "V": 32.4, "KA": 1.23},
            "omega_cv_pct": {"CL": 28.5, "V": 19.3},
            "theta_rse_pct": {"CL": 8.2, "V": 11.4, "KA": 15.6},
            "omega_rse_pct": {"CL": 22.1, "V": 31.0},
            "shrinkage_pct": {"CL": 12.4, "V": 18.7},
            "sigma": {"prop": 0.087, "add": None},
            "covariate_effects": [
                {"param": "CL", "covariate": "CRCL", "kind": "power",
                 "description": "CL increased with CRCL (exponent 0.75; 38% change low→high)",
                 "rse_pct": 18.3}
            ],
        }
    })


def _scm_state(base: PharmState | None = None) -> PharmState:
    st = base or _nlme_state()
    return st.model_copy(update={
        "scm_results": {
            "status": "ok", "label": "oral_1cmt",
            "base_ofv": -234.5, "final_ofv": -248.7,
            "forward_p": 0.05, "backward_p": 0.01, "n_candidates": 4,
            "selected": [{"param": "CL", "covariate": "CRCL",
                          "kind": "power", "delta_ofv": 14.2}],
            "steps": [{"phase": "forward", "effect": "CL~CRCL",
                       "delta_ofv": 14.2, "crit": 3.84, "df": 1,
                       "decision": "included"}],
        }
    })


# ── unit tests ────────────────────────────────────────────────────────────────

def test_empty_state_generates_file(tmp_path):
    result = generate_272(PharmState(), _ctx(tmp_path), {})
    p = Path(result.result["report_path"])
    assert p.exists() and p.suffix == ".docx"


def test_placeholders_when_no_analyses(tmp_path):
    result = generate_272(PharmState(), _ctx(tmp_path), {})
    doc = Document(result.result["report_path"])
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "To be completed" in full_text or "not yet" in full_text.lower()


def test_nca_state_contains_table_header(tmp_path):
    result = generate_272(_nca_state(), _ctx(tmp_path), {})
    doc = Document(result.result["report_path"])
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Table 2.7.2.1 header is in paragraph text for headings
    assert "2.7.2.1" in full_text
    assert "NCA" in full_text or "non-compartmental" in full_text.lower()


def test_nlme_state_table_272_populated(tmp_path):
    result = generate_272(_nlme_state(), _ctx(tmp_path), {})
    doc = Document(result.result["report_path"])
    # Table 2.7.2.2 should have theta rows; CL appears in table
    all_cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    assert any("CL" in c for c in all_cells)
    assert any("2.71" in c for c in all_cells)


def test_nlme_conclusions_mention_model(tmp_path):
    result = generate_272(_nlme_state(), _ctx(tmp_path), {"drug_name": "DrugX"})
    doc = Document(result.result["report_path"])
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "oral_1cmt" in full_text
    assert "converged" in full_text.lower()


def test_scm_selected_appears_in_special_pops(tmp_path):
    result = generate_272(_scm_state(), _ctx(tmp_path), {})
    doc = Document(result.result["report_path"])
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "CRCL" in full_text or "renal" in full_text.lower()


def test_study_info_args_propagate(tmp_path):
    result = generate_272(PharmState(), _ctx(tmp_path), {
        "drug_name": "TestMol", "sponsor": "ACME Pharma",
        "indication": "hypertension", "route": "oral",
    })
    doc = Document(result.result["report_path"])
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "TestMol" in full_text
    assert "ACME Pharma" in full_text
    assert "hypertension" in full_text


def test_reproducibility_section_present(tmp_path):
    result = generate_272(_base_state(), _ctx(tmp_path), {})
    doc = Document(result.result["report_path"])
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "PharmAgent" in full_text
    assert "SHA-256" in full_text
    # dataset sha should appear
    assert "abc123" in full_text


def test_writes_study_info_and_report_path(tmp_path):
    result = generate_272(PharmState(), _ctx(tmp_path), {
        "drug_name": "DrugY", "sponsor": "LabCo"
    })
    si = result.writes["study_info"]
    assert si.drug_name == "DrugY"              # StudyInfo object
    assert result.writes["regulatory_report_path"].endswith(".docx")
    assert result.result["study_info"]["drug_name"] == "DrugY"  # JSON-safe dict in result


def test_state_study_info_merges_with_args(tmp_path):
    """Args override state.study_info for the same key; unchanged keys from state persist."""
    st = PharmState(study_info=StudyInfo(drug_name="OldName", sponsor="OldSponsor"))
    result = generate_272(st, _ctx(tmp_path), {"drug_name": "NewName"})
    si = result.writes["study_info"]
    assert si.drug_name == "NewName"
    assert si.sponsor == "OldSponsor"


def test_full_state_no_crash(tmp_path):
    st = _scm_state()
    st = st.model_copy(update={"dose_prop_results": {"status": "ok",
        "proportional": True,
        "parameters": {"AUCinf": {"slope": 1.02, "slope_ci_lower": 0.95,
                                   "slope_ci_upper": 1.09, "proportional": True}}
    }})
    result = generate_272(st, _ctx(tmp_path), {"drug_name": "FullDrug"})
    p = Path(result.result["report_path"])
    assert p.exists()
    doc = Document(str(p))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Dose Proportionality" in full_text
    assert "FullDrug" in full_text


# ── HTTP endpoint tests ───────────────────────────────────────────────────────

def test_http_post_272_returns_report_path():
    from app.main import app
    client = TestClient(app)
    sid = client.post("/api/sessions").json()["id"]
    resp = client.post(f"/api/sessions/{sid}/report/272",
                       json={"drug_name": "DrugHTTP", "sponsor": "TestSponsor"})
    assert resp.status_code == 200
    data = resp.json()
    # ToolResult serialised into the chat-style response; report_path in result
    assert "result" in data or "report_path" in str(data)


def test_http_272_state_updated():
    from app.main import app
    client = TestClient(app)
    sid = client.post("/api/sessions").json()["id"]
    client.post(f"/api/sessions/{sid}/report/272",
                json={"drug_name": "StateDrug"})
    state = client.get(f"/api/sessions/{sid}/state").json()
    assert state["regulatory_report_path"] is not None
    assert state["study_info"]["drug_name"] == "StateDrug"


def test_http_272_download():
    from app.main import app
    client = TestClient(app)
    sid = client.post("/api/sessions").json()["id"]
    gen = client.post(f"/api/sessions/{sid}/report/272", json={})
    # extract filename from report path
    state = client.get(f"/api/sessions/{sid}/state").json()
    filename = Path(state["regulatory_report_path"]).name
    dl = client.get(f"/api/sessions/{sid}/report/272/{filename}")
    assert dl.status_code == 200
    assert dl.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument")


def test_http_272_wrong_filename_404():
    from app.main import app
    client = TestClient(app)
    sid = client.post("/api/sessions").json()["id"]
    client.post(f"/api/sessions/{sid}/report/272", json={})
    dl = client.get(f"/api/sessions/{sid}/report/272/nonexistent.docx")
    assert dl.status_code == 404
