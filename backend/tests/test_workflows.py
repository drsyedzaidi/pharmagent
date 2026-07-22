"""Workflow template integrity, and the report sections the popPK story ends in.

These tests are structural on purpose: they validate that every template step
names a real tool owned by a real agent, and that `poppk_full` orders its steps
so each tool's prerequisites are already in state. They deliberately do NOT run
`poppk_full` end to end — its NLME/SCM leg submits real population fits.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

from app.agents.definitions import AGENTS
from app.core.pharmstate import PharmState
from app.tools.base import ToolContext
from app.tools.builtins import default_registry
from app.tools.report_tools import generate_report
from app.workflows import WORKFLOWS, get_workflow


def _tools(name: str) -> list[str]:
    return [s["tool"] for s in get_workflow(name)["steps"]]


# ── template integrity (all workflows) ────────────────────────────────────────

def test_every_step_names_a_registered_tool_owned_by_its_agent():
    reg = default_registry()
    for wf_name, wf in WORKFLOWS.items():
        assert wf["name"] == wf_name, f"{wf_name}: name field disagrees with its key"
        assert wf["steps"], f"{wf_name}: no steps"
        for i, step in enumerate(wf["steps"]):
            tool = reg.get(step["tool"])  # raises KeyError if unregistered
            assert step["agent"] in AGENTS, \
                f"{wf_name} step {i}: agent {step['agent']!r} is not a defined agent"
            assert step["agent"] == tool.agent, \
                (f"{wf_name} step {i}: {step['tool']} is owned by {tool.agent!r}, "
                 f"but the step assigns it to {step['agent']!r}")
            assert step.get("label"), f"{wf_name} step {i}: missing label"


def test_every_workflow_loads_data_first_and_carries_a_human_gate():
    for wf_name, wf in WORKFLOWS.items():
        steps = wf["steps"]
        assert steps[0]["tool"] == "load_dataset", f"{wf_name}: must start by loading data"
        # This is a regulated-style app: no template may run start to finish
        # without a scientific decision point a human has to sign. A gate on the
        # final step is allowed — poppk_modeling deliberately ends at one.
        assert any(s.get("gate") for s in steps), f"{wf_name}: no human review gate"


# ── poppk_full ────────────────────────────────────────────────────────────────

def test_poppk_full_registered():
    wf = get_workflow("poppk_full")
    assert wf["name"] == "poppk_full"


def test_poppk_full_orders_the_nlme_leg_by_prerequisite():
    t = _tools("poppk_full")
    # run_nlme resolves its structural model from the last fit in state.
    assert t.index("fit_pk_model") < t.index("run_nlme")
    # run_diagnostics and run_vpc calibrate off a converged NLME fit of the same
    # model; without it they fall back to (or skip to) the two-stage path.
    assert t.index("run_nlme") < t.index("run_diagnostics")
    assert t.index("run_nlme") < t.index("run_vpc")
    # The forest needs covariate_effects from the NLME fit or the SCM result.
    assert t.index("run_nlme") < t.index("run_covariate_forest")
    assert t.index("run_scm") < t.index("run_covariate_forest")
    # The report is the last word.
    assert t[-1] == "generate_report"


def test_poppk_full_gates_before_the_expensive_fits():
    steps = get_workflow("poppk_full")["steps"]
    gated = [s["tool"] for s in steps if s.get("gate")]
    assert gated == ["fit_pk_model", "adversarial_review"]
    # The structural-model gate must precede every real population fit, so a
    # human confirms the model before the expensive leg commits.
    gate_i = next(i for i, s in enumerate(steps) if s["tool"] == "fit_pk_model")
    for expensive in ("run_nlme", "run_scm"):
        assert gate_i < next(i for i, s in enumerate(steps) if s["tool"] == expensive)


def test_poppk_full_excludes_nca_qc_and_unbounded_cost_tools():
    t = _tools("poppk_full")
    # run_qc scores an NCA analysis, not a mixed-effects fit (same rationale
    # poppk_modeling documents).
    assert "run_qc" not in t
    # run_simest needs a study `design` the template cannot know, and costs
    # several extra real NLME fits; run_engine_comparison belongs to
    # poppk_modeling.
    assert "run_simest" not in t
    assert "run_engine_comparison" not in t


def test_poppk_full_does_not_pin_the_auto_nlme_method():
    """`method="auto"` escalates to multiple seeded starts (hours on a real
    dataset) and is opt-in / job-backed — a template must not force it."""
    for step in get_workflow("poppk_full")["steps"]:
        assert "method" not in step.get("args", {})


def test_poppk_full_runs_to_the_structural_gate_and_stops(tmp_path):
    """The structural gate must actually hold at runtime: starting the workflow
    reaches it in seconds and executes NO population fit until a human resumes.
    """
    import itertools

    from app.core.llm import MockLLM
    from app.core.orchestrator import Orchestrator
    from app.core.store import SessionStore

    sample = str(Path(__file__).parent.parent / "sample_data" / "oral_pk.csv")
    orch = Orchestrator(llm=MockLLM(), clock=lambda c=itertools.count(): f"t{next(c)}",
                        store=SessionStore(":memory:"))
    sid = orch.create_session().id
    out = orch.start_workflow(sid, "poppk_full", {"path": sample})

    assert out["status"] == "awaiting_review"
    assert out["review"]["label"] == "Compare structural models"
    ran = [e["tool"] for e in out["executed"]]
    assert ran == ["load_dataset", "profile_pk_dataset", "validate_cdisc",
                   "generate_spaghetti_plot", "fit_pk_model"]
    # The expensive leg is untouched — this is the whole point of the gate.
    assert out["state"]["nlme_results"] is None
    assert out["state"]["scm_results"] is None
    assert out["audit_ok"]

    # Rejecting at the gate must also leave the population fits unrun.
    rej = orch.resume_workflow(sid, approve=False, actor="alice", reason="wrong model")
    assert rej["status"] == "rejected"
    assert rej["state"]["nlme_results"] is None


# ── the report end of the story ───────────────────────────────────────────────

def _state_with_diagnostics_and_forest() -> PharmState:
    return PharmState(
        dataset_metadata={"dataset_id": "ds", "n_subjects": 12},
        diagnostics_results={
            "status": "ok", "model_key": "onecmt_oral", "label": "1-cmt oral",
            "nlme_provenance": "onecmt_oral",
            "residuals": {"summary": {"n": 120, "mean": 0.01, "sd": 0.98}},
            "cwres": {"status": "ok",
                      "summary": {"cwres_mean": -0.02, "cwres_sd": 1.03, "n": 120}},
            "npde": {"status": "ok",
                     "summary": {"mean": 0.04, "sd": 1.01, "n": 120,
                                 "pct_outside_1_96": 5.8}},
        },
        forest_results={
            "status": "ok", "model_key": "onecmt_oral", "label": "1-cmt oral",
            "source": "scm", "percentiles": [5.0, 95.0], "ci_level": 0.90,
            "x_range": [0.6, 1.5], "bounds": None,
            "summary": {"n_rows": 2, "n_effects": 1},
            "notes": ["A covariate acts on a weight-scaled parameter."],
            "rows": [
                {"param": "CL", "covariate": "WT", "kind": "continuous",
                 "eval_label": "WT=55", "eval_value": 55.0, "gmr": 0.82,
                 "ci_lo": 0.74, "ci_hi": 0.91, "ci_source": "delta",
                 "omega_cv_pct": 28.0, "allometric_note": True},
                {"param": "CL", "covariate": "WT", "kind": "continuous",
                 "eval_label": "WT=95", "eval_value": 95.0, "gmr": 1.19,
                 "ci_lo": None, "ci_hi": None, "ci_source": "unavailable",
                 "omega_cv_pct": 28.0, "allometric_note": True},
            ],
        },
    )


def _report_text(state: PharmState, tmp_path: Path) -> str:
    ctx = ToolContext(data_dir=str(tmp_path))
    res = generate_report(state, ctx, {})
    doc = Document(res.result["report_path"])
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        parts += [c.text for row in table.rows for c in row.cells]
    return "\n".join(parts)


def test_report_carries_the_diagnostics_and_forest_sections(tmp_path):
    text = _report_text(_state_with_diagnostics_and_forest(), tmp_path)
    assert "Residual diagnostics" in text
    assert "CWRES" in text and "npd" in text
    assert "Covariate effects (forest)" in text
    assert "WT=55" in text
    # The unavailable-CI row must still appear, with its CI rendered as missing.
    assert "WT=95" in text
    # Caveats travel with the numbers.
    assert "weight-scaled parameter" in text


def test_report_omits_those_sections_when_the_tools_did_not_run(tmp_path):
    text = _report_text(PharmState(dataset_metadata={"dataset_id": "ds"}), tmp_path)
    assert "Residual diagnostics" not in text
    assert "Covariate effects (forest)" not in text


def test_report_skips_diagnostics_that_reported_a_non_ok_status(tmp_path):
    state = PharmState(
        dataset_metadata={"dataset_id": "ds"},
        diagnostics_results={"status": "no_fit", "message": "Fit a PK model first."},
        forest_results={"status": "no_fit", "message": "No covariate model."},
    )
    text = _report_text(state, tmp_path)
    assert "Residual diagnostics" not in text
    assert "Covariate effects (forest)" not in text
