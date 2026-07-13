"""Report Agent tools: generate a regulatory-style DOCX from PharmState."""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from app.core.pharmstate import PharmState
from app.tools.base import Tool, ToolContext, ToolResult

METHODS_NCA = (
    "Non-compartmental analysis was performed using the linear-up/log-down "
    "trapezoidal rule. The terminal elimination rate constant (lambda_z) was "
    "estimated by log-linear regression over the terminal phase, selecting the "
    "window that maximized the adjusted R-squared. AUC was extrapolated to "
    "infinity as Clast/lambda_z. Clearance (CL/F) and volume (Vz/F) were "
    "derived from dose and AUCinf. All computations were performed by "
    "deterministic numerical routines; results are recorded in a SHA-256 "
    "hash-chained audit trail."
)


def _fmt(v: Any, d: int = 2) -> str:
    """Format a numeric cell to d decimals; '-' for missing."""
    if v is None:
        return "-"
    try:
        return f"{float(v):.{d}f}"
    except (TypeError, ValueError):
        return str(v)


def _subject_key(r: dict[str, Any]) -> tuple[int, Any]:
    """Sort subjects numerically when possible, else lexically."""
    s = r.get("subject")
    try:
        return (0, int(float(s)))
    except (TypeError, ValueError):
        return (1, str(s))


# (display header, nca_parameters key, decimals)
_SUBJECT_COLS: list[tuple[str, str, int]] = [
    ("ID", "subject", 0),
    ("Dose", "dose", 0),
    ("Cmax", "Cmax", 2),
    ("Tmax", "Tmax", 2),
    ("AUClast", "AUC_last", 1),
    ("AUCinf", "AUC_inf", 1),
    ("t½", "t_half", 1),
    ("CL/F", "CL_F", 2),
    ("Vz/F", "Vz_F", 1),
    ("λz adjR²", "lambda_z_r2_adj", 3),
    ("λz n", "lambda_z_n_points", 0),
    ("%extrap", "pct_AUC_extrap", 1),
]


def _subject_table(doc: Document, params: list[dict[str, Any]]) -> None:
    """Primary results: one row per subject. Robust for weight-based dosing
    where each subject has a near-unique dose (no meaningful pooling)."""
    if not params:
        doc.add_paragraph("No per-subject NCA parameters available.")
        return
    rows = sorted(params, key=_subject_key)
    headers = [h for h, _, _ in _SUBJECT_COLS]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r in rows:
        cells = table.add_row().cells
        for i, (_h, key, dec) in enumerate(_SUBJECT_COLS):
            cells[i].text = _fmt(r.get(key), dec)


def _dose_grouped_meaningfully(summary: dict[str, Any]) -> bool:
    """Only pool by dose when there are genuinely few levels with replicates —
    otherwise (e.g. mg/kg weight-based dosing) every subject is its own 'group'
    and the geomean/CV table is noise. Mirrors the frontend heuristic."""
    rows = summary.get("by_dose", [])
    n_subjects = summary.get("n_subjects", 0) or 0
    if not rows or not n_subjects:
        return False
    few_levels = len(rows) <= max(3, n_subjects // 3)
    has_replicates = any((r.get("n") or 0) >= 3 for r in rows)
    return few_levels and has_replicates


def _summary_table(doc: Document, summary: dict[str, Any]) -> None:
    rows = summary.get("by_dose", [])
    if not rows:
        return
    cols = ["dose", "n", "Cmax_geomean", "Cmax_geocv_pct",
            "AUC_inf_geomean", "AUC_inf_geocv_pct", "t_half_median"]
    headers = ["Dose", "n", "Cmax GM", "Cmax gCV%", "AUCinf GM", "AUCinf gCV%", "t½ median"]
    decimals = {"dose": 0, "n": 0, "Cmax_geomean": 2, "Cmax_geocv_pct": 1,
                "AUC_inf_geomean": 1, "AUC_inf_geocv_pct": 1, "t_half_median": 1}
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r in rows:
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            cells[i].text = _fmt(r.get(c), decimals[c])


def _kv_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = _fmt(v) if isinstance(v, (int, float)) else str(v)


def _extra_sections(doc: Document, state: PharmState) -> None:
    """Append sections for any further analyses present in state."""
    n = 4

    be = state.be_results
    if be and be.get("status") == "ok":
        n += 1
        doc.add_heading(f"{n}. Bioequivalence", level=1)
        doc.add_paragraph(f"{be['design']}: {be.get('test_level')} vs {be.get('reference_level')}; "
                          f"limits {be['limits'][0]}–{be['limits'][1]}%. Overall: "
                          f"{'BIOEQUIVALENT' if be.get('bioequivalent') else 'NOT bioequivalent'}.")
        _kv_table(doc, ["Parameter", "GMR %", "90% CI", "Within"],
                  [[k, v.get("gmr_pct"), f"{_fmt(v.get('ci_lower_pct'))}–{_fmt(v.get('ci_upper_pct'))}",
                    "yes" if v.get("within_limits") else "no"] for k, v in (be.get("parameters") or {}).items()])

    dp = state.dose_prop_results
    if dp and dp.get("status") == "ok":
        n += 1
        doc.add_heading(f"{n}. Dose proportionality", level=1)
        doc.add_paragraph(f"Power model; {'proportional' if dp.get('proportional') else 'not proportional'}.")
        _kv_table(doc, ["Parameter", "Slope", "90% CI", "Proportional"],
                  [[k, v.get("slope"), f"{_fmt(v.get('slope_ci_lower'))}–{_fmt(v.get('slope_ci_upper'))}",
                    "yes" if v.get("proportional") else "no"] for k, v in (dp.get("parameters") or {}).items()])

    pm = state.pk_model_results
    if pm and pm.get("status") == "ok":
        n += 1
        best = pm if pm.get("mode") == "fit" else (pm.get("best") or {})
        doc.add_heading(f"{n}. Structural PK model", level=1)
        doc.add_paragraph(f"Model: {best.get('label')} — {best.get('n_converged')}/{best.get('n_subjects')} "
                          f"converged, mean AIC {best.get('mean_aic')}. Two-stage population:")
        pop = (best.get("population") or {}).get("parameters") or {}
        _kv_table(doc, ["Parameter", "Typical (GM)", "IIV CV%", "n"],
                  [[k, v.get("typical_value"), v.get("iiv_cv_pct"), v.get("n")] for k, v in pop.items()])

    pk = state.poppk_results
    if pk and pk.get("status") == "ok":
        n += 1
        doc.add_heading(f"{n}. Population PK (two-stage)", level=1)
        _kv_table(doc, ["Parameter", "Typical (GM)", "IIV CV%", "n"],
                  [[k, v.get("typical_value"), v.get("iiv_cv_pct"), v.get("n")]
                   for k, v in (pk.get("parameters") or {}).items()])

    nl = state.nlme_results
    if nl and nl.get("status") == "ok":
        n += 1
        doc.add_heading(f"{n}. Population PK — mixed-effects ({nl.get('method')})", level=1)
        cond = nl.get("condition_number")
        cond_str = f"; condition number {_fmt(cond, 1)}" if cond is not None else ""
        doc.add_paragraph(f"Model: {nl.get('label')}; IIV on {nl.get('iiv_params')}; "
                          f"{nl.get('error_model')} error; OFV {_fmt(nl.get('ofv'), 1)}{cond_str}; "
                          f"{'converged' if nl.get('converged') else 'did not converge'}.")
        omega, shr = nl.get("omega_cv_pct") or {}, nl.get("shrinkage_pct") or {}
        rse, ormse = nl.get("theta_rse_pct") or {}, nl.get("omega_rse_pct") or {}
        _kv_table(doc, ["Parameter", "Typical (θ)", "RSE%", "IIV CV%", "IIV RSE%", "η-shrinkage%"],
                  [[p, v, rse.get(p, "—"), omega.get(p, "—"), ormse.get(p, "—"), shr.get(p, "—")]
                   for p, v in (nl.get("theta") or {}).items()])
        ce = nl.get("covariate_effects") or []
        if ce:
            doc.add_paragraph("Covariate effects: " + "; ".join(
                f"{e['param']} ~ {e['description']}"
                + (f" ({_fmt(e['rse_pct'], 1)}% RSE)"
                   if isinstance(e.get("rse_pct"), (int, float)) else "")
                for e in ce))
        if nl.get("cov_note"):
            doc.add_paragraph(str(nl.get("cov_note")))

    sc = state.scm_results
    if sc and sc.get("status") == "ok":
        n += 1
        doc.add_heading(f"{n}. Covariate model selection (SCM)", level=1)
        sel = sc.get("selected") or []
        sel_str = ", ".join(f"{s['param']}~{s['covariate']} ({s['kind']})"
                            for s in sel) or "none"
        doc.add_paragraph(
            f"Stepwise covariate modeling on {sc.get('label')} "
            f"(forward p<{sc.get('forward_p')}, backward p<{sc.get('backward_p')}): "
            f"{sc.get('n_candidates')} candidate(s) tested; selected {sel_str}; "
            f"OFV {_fmt(sc.get('base_ofv'), 1)} -> {_fmt(sc.get('final_ofv'), 1)}.")
        steps = sc.get("steps") or []
        if steps:
            _kv_table(doc, ["Phase", "Effect", "ΔOFV", "χ² crit", "df", "Decision"],
                      [[s["phase"], s["effect"], _fmt(s["delta_ofv"], 2),
                        _fmt(s["crit"], 2), s["df"], s["decision"]] for s in steps])

    vp = state.vpc_results
    if vp and vp.get("status") == "ok":
        n += 1
        g = vp.get("gof") or {}
        doc.add_heading(f"{n}. Model evaluation (GOF / VPC)", level=1)
        doc.add_paragraph(f"Goodness-of-fit (log scale): R²(IPRED) = {_fmt(g.get('r2_log_ipred'), 3)}, "
                          f"RMSE = {_fmt(g.get('rmse_log_ipred'), 3)}, n = {g.get('n')}. "
                          f"A 5/50/95 VPC band was generated at the {vp.get('vpc_dose')} dose level.")


def _reproducibility_section(doc: Document, state: PharmState) -> None:
    """Software/platform provenance, dataset integrity hash, and seed — so the
    report's results can be reproduced from the record alone (ALCOA+)."""
    from app.core.provenance import collect_provenance
    prov = collect_provenance()
    meta = state.dataset_metadata or {}
    doc.add_heading("Software & reproducibility", level=1)
    doc.add_paragraph(
        f"Generated by PharmAgent v{prov['app_version']} (git {prov['git_sha']}) on "
        f"Python {prov['python']}, {prov['platform']}. Numerical libraries: "
        f"numpy {prov['numpy']}, scipy {prov['scipy']}, pandas {prov['pandas']}.")
    sha = meta.get("dataset_sha256")
    if sha and sha != "n/a":
        doc.add_paragraph(f"Source dataset SHA-256: {sha}")
    doc.add_paragraph(
        "Stochastic procedures (SAEM E-step, VPC/pcVPC, simulation-based NPDE) use "
        "a fixed RNG seed (20250614) so results are bit-reproducible; FOCE-I is "
        "deterministic. All computations are recorded in a SHA-256 hash-chained, "
        "identity-stamped audit trail (verify via GET /sessions/{id}/audit).")


def generate_report(state: PharmState, ctx: ToolContext, args: dict[str, Any]) -> ToolResult:
    title = args.get("title", "Pharmacokinetic Analysis Report")
    doc = Document()
    doc.add_heading(title, level=0)

    meta = state.dataset_metadata or {}
    doc.add_heading("1. Dataset", level=1)
    doc.add_paragraph(
        f"Dataset {meta.get('dataset_id', 'n/a')}: {meta.get('n_records', '?')} records, "
        f"{meta.get('n_subjects', '?')} subjects, dose levels "
        f"{meta.get('dose_levels', 'n/a')}.")

    doc.add_heading("2. Methods", level=1)
    doc.add_paragraph(METHODS_NCA)

    doc.add_heading("3. Results", level=1)
    params = state.nca_parameters or []
    doc.add_heading("3.1 Per-subject NCA parameters", level=2)
    doc.add_paragraph(
        f"Non-compartmental parameters for all {len(params)} subjects. "
        "λz adjR² and λz n report the terminal-phase fit quality and the number "
        "of points used; %extrap is the fraction of AUCinf obtained by "
        "extrapolation beyond the last measured concentration.")
    _subject_table(doc, params)

    if state.nca_summary and _dose_grouped_meaningfully(state.nca_summary):
        doc.add_heading("3.2 Dose-group summary", level=2)
        doc.add_paragraph(
            "Geometric mean (GM) and geometric CV% by dose level; t½ as median.")
        _summary_table(doc, state.nca_summary)

    doc.add_heading("4. Quality Control", level=1)
    doc.add_paragraph(f"QC verdict: {state.qc_verdict or 'not run'}.")
    for c in (state.qc_checklist or []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{c['status']}] {c['check']}: ").bold = True
        p.add_run(c["detail"])

    _extra_sections(doc, state)
    _reproducibility_section(doc, state)

    out_dir = Path(ctx.data_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    path = str(out_dir / f"report_{meta.get('dataset_id', 'analysis')}.docx")
    doc.save(path)

    sections = {
        "dataset": meta.get("dataset_id"),
        "methods": "NCA (linear-up/log-down, best-fit lambda_z)",
        "qc_verdict": state.qc_verdict,
    }
    return ToolResult(
        summary=f"Generated DOCX report at {path}.",
        action=f"generate_report -> {path}",
        writes={"report_path": path, "report_sections": sections},
        result={"report_path": path, "sections": sections},
    )


TOOLS = [
    Tool("generate_report",
         "Generate an FDA/EMA-style DOCX report (dataset, methods, per-subject "
         "NCA table, optional dose-group summary, QC) from the current state.",
         "report",
         {"type": "object",
          "properties": {"title": {"type": "string"}}, "required": []},
         generate_report),
]
