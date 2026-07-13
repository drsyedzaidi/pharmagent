"""ICH M4E Module 2.7.2 — Summary of Clinical Pharmacology Studies.

Generates a CTD-structured DOCX from whatever PharmState analysis results exist.
Empty sections are rendered as explicit placeholders so the writer knows what
still needs manual input before submission.

Table numbering follows ICH M4E convention:
  Table 2.7.2.1   NCA summary by dose group
  Table 2.7.2.2   Population PK parameter estimates (theta / RSE / IIV)
  Table 2.7.2.3   Covariate effects on PK parameters
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from app.core.pharmstate import PharmState, StudyInfo
from app.tools.base import Tool, ToolContext, ToolResult

# ── helpers ───────────────────────────────────────────────────────────────────

def _fmt(v: Any, d: int = 2) -> str:
    if v is None:
        return "—"
    try:
        return f"{float(v):.{d}f}"
    except (TypeError, ValueError):
        return str(v)


def _placeholder(doc: Document, text: str = "[To be completed]") -> None:
    p = doc.add_paragraph(text)
    p.runs[0].italic = True


def _kv_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    if not rows:
        _placeholder(doc, "[No data available]")
        return
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for r in rows:
        cells = tbl.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = _fmt(v) if isinstance(v, (int, float)) else str(v or "—")


# ── section renderers ─────────────────────────────────────────────────────────

def _section_21_background(doc: Document, si: StudyInfo, meta: dict) -> None:
    drug = si.drug_name or "[drug name]"
    route = si.route or "oral"
    indication = si.indication or "[indication]"
    pop = si.pop_description or "healthy adult volunteers"
    dose_range = si.dose_range or (meta.get("dose_levels") or "[dose range]")
    matrix = si.matrix or "plasma"
    lloq = si.assay_lloq or "[LLOQ]"
    sponsor = si.sponsor or "[sponsor]"
    study_id = si.study_id or "[study ID]"

    doc.add_heading("2.7.2.1  Background", level=1)
    doc.add_paragraph(
        f"{drug} is a {route} agent under investigation for the treatment of "
        f"{indication}. This module summarises the clinical pharmacology programme "
        f"conducted by {sponsor} (Study {study_id})."
    )
    doc.add_paragraph(
        f"The pharmacokinetic (PK) profile of {drug} was characterised in "
        f"{pop} following {route} administration of doses ranging from {dose_range}. "
        f"{drug} concentrations were quantified in {matrix} using a validated "
        f"bioanalytical method with a lower limit of quantification (LLOQ) of {lloq}."
    )


def _section_22_nca(doc: Document, state: PharmState, si: StudyInfo) -> None:
    doc.add_heading("2.7.2.2  Summary of Results of Individual Studies", level=1)
    doc.add_heading("2.7.2.2.1  Healthy Subject / Single-Dose PK Studies", level=2)

    meta = state.dataset_metadata or {}
    n_subj = meta.get("n_subjects", "?")
    pop = si.pop_description or "healthy adult volunteers"
    drug = si.drug_name or "[drug]"

    if state.nca_parameters:
        doc.add_paragraph(
            f"Non-compartmental analysis (NCA) was applied to characterise the "
            f"pharmacokinetics of {drug} in {n_subj} {pop}. The linear-up/log-down "
            f"trapezoidal rule was used to estimate AUC; the terminal elimination "
            f"rate constant (λz) was estimated by log-linear regression over the "
            f"terminal phase. Key PK parameters are summarised in Table 2.7.2.1."
        )
    else:
        _placeholder(doc, "[NCA results not yet available — run NCA analysis first]")


def _section_22_poppk(doc: Document, state: PharmState) -> None:
    doc.add_heading("2.7.2.2.2  Population PK Analysis", level=2)
    nl = state.nlme_results
    if not nl or nl.get("status") != "ok":
        _placeholder(doc, "[Population PK analysis not yet performed]")
        return

    method = (nl.get("method") or "FOCE-I").upper()
    label = nl.get("label", "[model]")
    iiv = nl.get("iiv_params") or []
    err = nl.get("error_model", "proportional")
    ofv = nl.get("ofv")
    converged = "converged successfully" if nl.get("converged") else "did not converge"

    doc.add_paragraph(
        f"A population pharmacokinetic (PopPK) model was developed using nonlinear "
        f"mixed-effects (NLME) methodology ({method}, PharmAgent v0.1.0). A {label} "
        f"structural model with inter-individual variability (IIV) on "
        f"{', '.join(iiv) if iiv else 'selected parameters'} and a {err} residual "
        f"error model best described the concentration–time data (OFV {_fmt(ofv, 1)}; "
        f"{converged}). Final model parameter estimates are presented in Table 2.7.2.2."
    )

    cond = nl.get("condition_number")
    if cond is not None:
        stability = ("satisfactory model stability"
                     if cond < 1000 else
                     "potential numerical instability; interpret standard errors with caution")
        doc.add_paragraph(
            f"The condition number of the covariance matrix was {_fmt(cond, 1)}, "
            f"indicating {stability}."
        )

    vpc = state.vpc_results
    if vpc and vpc.get("status") == "ok":
        g = vpc.get("gof") or {}
        doc.add_paragraph(
            f"Model adequacy was assessed by standard goodness-of-fit (GOF) diagnostics. "
            f"The coefficient of determination between observed and individual predicted "
            f"concentrations (log-scale) was R² = {_fmt(g.get('r2_log_ipred'), 3)}. "
            f"A visual predictive check (VPC) was performed to evaluate the model's "
            f"ability to reproduce the observed distribution of concentrations."
        )


def _section_23_analyses(doc: Document, state: PharmState) -> None:
    doc.add_heading("2.7.2.3  Comparison and Analysis Across Studies", level=1)

    dp = state.dose_prop_results
    if dp and dp.get("status") == "ok":
        doc.add_heading("2.7.2.3.1  Dose Proportionality", level=2)
        prop = "proportional" if dp.get("proportional") else "not proportional"
        doc.add_paragraph(
            f"Dose proportionality was assessed using the power model. Exposure was "
            f"{prop} over the dose range evaluated."
        )
        params = dp.get("parameters") or {}
        _kv_table(doc, ["Parameter", "Slope", "90% CI", "Proportional?"],
                  [[k, v.get("slope"), f"{_fmt(v.get('slope_ci_lower'))}–{_fmt(v.get('slope_ci_upper'))}",
                    "Yes" if v.get("proportional") else "No"] for k, v in params.items()])
    else:
        doc.add_heading("2.7.2.3.1  Dose Proportionality", level=2)
        _placeholder(doc, "[Dose proportionality analysis not performed]")

    be = state.be_results
    if be and be.get("status") == "ok":
        doc.add_heading("2.7.2.3.2  Bioequivalence", level=2)
        overall = "BIOEQUIVALENT" if be.get("bioequivalent") else "NOT bioequivalent"
        doc.add_paragraph(
            f"Bioequivalence study design: {be.get('design')} "
            f"({be.get('test_level')} vs {be.get('reference_level')}; "
            f"acceptance limits {be['limits'][0]}–{be['limits'][1]}%). "
            f"Overall conclusion: {overall}."
        )


def _section_24_special_pops(doc: Document, state: PharmState) -> None:
    doc.add_heading("2.7.2.4  Special Populations", level=1)
    nl = state.nlme_results
    ce = (nl.get("covariate_effects") or []) if nl and nl.get("status") == "ok" else []
    sc = state.scm_results
    sel = (sc.get("selected") or []) if sc and sc.get("status") == "ok" else []

    if not ce and not sel:
        _placeholder(doc,
            "[No covariate analysis results available. Run SCM or include covariates "
            "in the NLME model to populate this section.]")
        return

    doc.add_paragraph(
        "The influence of intrinsic and extrinsic factors on the PK of the drug was "
        "evaluated in the population PK model. Covariate effects are summarised in "
        "Table 2.7.2.3."
    )

    # Build a narrative from the covariate effects list in nlme_results
    cov_by_type: dict[str, list[str]] = {
        "renal": [], "hepatic": [], "age": [], "weight": [],
        "sex": [], "other": [],
    }
    for e in ce:
        cov = (e.get("covariate") or "").lower()
        desc = e.get("description") or str(e)
        param = e.get("param", "")
        narrative = f"{param}: {desc}"
        if any(k in cov for k in ("crcl", "egfr", "renal", "gfr")):
            cov_by_type["renal"].append(narrative)
        elif any(k in cov for k in ("alt", "ast", "bilirubin", "hepat", "child")):
            cov_by_type["hepatic"].append(narrative)
        elif cov in ("age",):
            cov_by_type["age"].append(narrative)
        elif cov in ("wt", "weight", "bmi"):
            cov_by_type["weight"].append(narrative)
        elif cov in ("sex", "gender"):
            cov_by_type["sex"].append(narrative)
        else:
            cov_by_type["other"].append(narrative)

    subsection_map = {
        "renal": "2.7.2.4.1  Renal Impairment",
        "hepatic": "2.7.2.4.2  Hepatic Impairment",
        "age": "2.7.2.4.3  Age",
        "weight": "2.7.2.4.4  Body Weight",
        "sex": "2.7.2.4.5  Sex",
        "other": "2.7.2.4.6  Other Covariates",
    }
    for key, title in subsection_map.items():
        items = cov_by_type[key]
        if items:
            doc.add_heading(title, level=2)
            for item in items:
                doc.add_paragraph(item)

    if not any(cov_by_type.values()):
        _placeholder(doc, "[Covariate details: see Table 2.7.2.3]")


def _section_25_ddi(doc: Document) -> None:
    doc.add_heading("2.7.2.5  Drug–Drug Interactions", level=1)
    _placeholder(doc,
        "[No dedicated DDI studies have been conducted to date. "
        "In vitro CYP inhibition/induction data and clinical DDI study results "
        "will be summarised here when available.]")


def _section_26_conclusions(doc: Document, state: PharmState, si: StudyInfo) -> None:
    doc.add_heading("2.7.2.6  Conclusions", level=1)
    drug = si.drug_name or "[drug]"
    parts: list[str] = []

    if state.nca_parameters:
        meta = state.dataset_metadata or {}
        parts.append(
            f"The pharmacokinetics of {drug} were characterised by NCA in "
            f"{meta.get('n_subjects', '?')} {si.pop_description or 'subjects'}."
        )

    nl = state.nlme_results
    if nl and nl.get("status") == "ok":
        label = nl.get("label", "a compartmental model")
        parts.append(
            f"A {label} adequately described the observed concentration–time "
            f"profiles across the dose range studied."
        )
        iiv = nl.get("iiv_params") or []
        if iiv:
            omega = nl.get("omega_cv_pct") or {}
            cv_str = ", ".join(
                f"{p} ({_fmt(omega.get(p), 0)}% CV)"
                for p in iiv if p in omega
            )
            if cv_str:
                parts.append(f"Inter-individual variability was estimated for {cv_str}.")

    sc = state.scm_results
    if sc and sc.get("status") == "ok":
        sel = sc.get("selected") or []
        if sel:
            effects = ", ".join(f"{s['param']}~{s['covariate']}" for s in sel)
            parts.append(
                f"Stepwise covariate modelling identified the following statistically "
                f"significant covariate relationships: {effects}."
            )
        else:
            parts.append("No statistically significant covariates were identified.")

    if not parts:
        _placeholder(doc,
            "[Complete pharmacokinetic analyses (NCA, population PK) before "
            "finalising the conclusions section.]")
        return

    for part in parts:
        doc.add_paragraph(part)


# ── table renderers ───────────────────────────────────────────────────────────

def _table_2721_nca(doc: Document, state: PharmState) -> None:
    """Table 2.7.2.1 — NCA summary statistics by dose group."""
    doc.add_heading("Table 2.7.2.1  Summary of NCA Parameters by Dose Group", level=2)
    summary = state.nca_summary or {}
    rows = summary.get("by_dose") or []
    if not rows:
        _placeholder(doc, "[NCA summary not available]")
        return
    _kv_table(doc,
              ["Dose", "n", "Cmax GM (CV%)", "AUCinf GM (CV%)", "t½ median", "CL/F GM"],
              [[r.get("dose"), r.get("n"),
                f"{_fmt(r.get('Cmax_geomean'))} ({_fmt(r.get('Cmax_geocv_pct'), 1)}%)",
                f"{_fmt(r.get('AUC_inf_geomean'), 1)} ({_fmt(r.get('AUC_inf_geocv_pct'), 1)}%)",
                _fmt(r.get("t_half_median"), 1),
                _fmt(r.get("CL_F_geomean"))] for r in rows])


def _table_2722_poppk(doc: Document, state: PharmState) -> None:
    """Table 2.7.2.2 — Final population PK parameter estimates."""
    doc.add_heading("Table 2.7.2.2  Final Population PK Parameter Estimates", level=2)
    nl = state.nlme_results
    if not nl or nl.get("status") != "ok":
        _placeholder(doc, "[Population PK not performed]")
        return
    theta = nl.get("theta") or {}
    omega = nl.get("omega_cv_pct") or {}
    rse = nl.get("theta_rse_pct") or {}
    ormse = nl.get("omega_rse_pct") or {}
    shr = nl.get("shrinkage_pct") or {}
    rows = [
        [p, _fmt(v), _fmt(rse.get(p)), _fmt(omega.get(p), 1),
         _fmt(ormse.get(p)), _fmt(shr.get(p), 1)]
        for p, v in theta.items()
    ]
    if not rows:
        _placeholder(doc, "[No parameter estimates available]")
        return
    _kv_table(doc,
              ["Parameter", "Estimate", "%RSE", "IIV CV%", "IIV %RSE", "η-shrinkage%"],
              rows)
    sigma = nl.get("sigma") or {}
    if sigma:
        sig_rows = [[k, _fmt(v)] for k, v in sigma.items() if v is not None]
        if sig_rows:
            doc.add_paragraph("Residual variability (σ):")
            _kv_table(doc, ["Component", "Estimate"], sig_rows)


def _table_2723_covariates(doc: Document, state: PharmState) -> None:
    """Table 2.7.2.3 — Covariate effects on PK parameters."""
    doc.add_heading("Table 2.7.2.3  Covariate Effects on PK Parameters", level=2)
    nl = state.nlme_results
    ce = (nl.get("covariate_effects") or []) if nl and nl.get("status") == "ok" else []
    sc = state.scm_results
    sel = (sc.get("selected") or []) if sc and sc.get("status") == "ok" else []

    rows: list[list[Any]] = []
    for e in ce:
        rse_str = f"{_fmt(e.get('rse_pct'), 1)}%" if e.get("rse_pct") is not None else "—"
        rows.append([e.get("param", "—"), e.get("covariate", "—"),
                     e.get("kind", "—"), e.get("description", "—"), rse_str])

    if not rows and sel:
        for s in sel:
            delta = _fmt(s.get("delta_ofv"), 1)
            rows.append([s.get("param", "—"), s.get("covariate", "—"),
                         s.get("kind", "—"), f"ΔOFV {delta}", "—"])

    if not rows:
        _placeholder(doc, "[No covariate effects estimated]")
        return
    _kv_table(doc, ["Parameter", "Covariate", "Model", "Description / Effect", "%RSE"],
              rows)


# ── main tool function ────────────────────────────────────────────────────────

def generate_272(state: PharmState, ctx: ToolContext, args: dict[str, Any]) -> ToolResult:
    """Generate an ICH M4E Module 2.7.2 CTD-compliant DOCX."""
    # Merge args into study info; existing study_info in state takes lower priority
    si_data = (state.study_info.model_dump() if state.study_info else {})
    for k in ("drug_name", "sponsor", "study_id", "route", "indication",
              "pop_description", "dose_range", "matrix", "assay_lloq"):
        if args.get(k):
            si_data[k] = args[k]
    si = StudyInfo(**si_data)

    meta = state.dataset_metadata or {}
    doc = Document()

    # ── cover ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("Module 2.7.2", level=0)
    title.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_heading("Summary of Clinical Pharmacology Studies", level=1)
    sub.alignment = 1
    if si.drug_name:
        p = doc.add_paragraph(f"Drug: {si.drug_name}")
        p.alignment = 1
    if si.sponsor:
        p = doc.add_paragraph(f"Sponsor: {si.sponsor}")
        p.alignment = 1
    if si.study_id:
        p = doc.add_paragraph(f"Study: {si.study_id}")
        p.alignment = 1
    from app.core.provenance import collect_provenance
    prov = collect_provenance()
    p = doc.add_paragraph(f"Generated by PharmAgent v{prov['app_version']}")
    p.alignment = 1
    doc.add_page_break()

    # ── confidentiality notice ─────────────────────────────────────────────
    notice = doc.add_paragraph(
        "CONFIDENTIAL — This document contains proprietary information of "
        f"{si.sponsor or '[Sponsor]'}. Do not copy or distribute without written permission."
    )
    notice.runs[0].font.size = Pt(9)
    notice.runs[0].italic = True
    doc.add_paragraph()

    # ── body sections ──────────────────────────────────────────────────────
    _section_21_background(doc, si, meta)
    _section_22_nca(doc, state, si)
    _section_22_poppk(doc, state)
    _section_23_analyses(doc, state)
    _section_24_special_pops(doc, state)
    _section_25_ddi(doc)
    _section_26_conclusions(doc, state, si)

    # ── tables ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Tables", level=1)
    _table_2721_nca(doc, state)
    doc.add_paragraph()
    _table_2722_poppk(doc, state)
    doc.add_paragraph()
    _table_2723_covariates(doc, state)

    # ── reproducibility ───────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Software and Reproducibility", level=1)
    sha = meta.get("dataset_sha256")
    doc.add_paragraph(
        f"Generated by PharmAgent v{prov['app_version']} (git {prov['git_sha']}) on "
        f"Python {prov['python']}, {prov['platform']}. "
        f"Numerical libraries: numpy {prov['numpy']}, scipy {prov['scipy']}, "
        f"pandas {prov['pandas']}."
    )
    if sha and sha != "n/a":
        doc.add_paragraph(f"Source dataset SHA-256: {sha}")
    doc.add_paragraph(
        "All computations are recorded in a SHA-256 hash-chained, identity-stamped "
        "audit trail (GET /sessions/{id}/audit). Stochastic procedures use a fixed "
        "RNG seed; FOCE-I is deterministic."
    )

    # ── write file ─────────────────────────────────────────────────────────
    out_dir = Path(ctx.data_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    study_id_slug = (si.study_id or meta.get("dataset_id") or "analysis").replace(" ", "_")
    path = str(out_dir / f"module_272_{study_id_slug}.docx")
    doc.save(path)

    return ToolResult(
        summary=f"Generated ICH Module 2.7.2 DOCX at {path}.",
        action=f"generate_272 -> {path}",
        writes={
            "study_info": si,            # StudyInfo object for PharmState
            "regulatory_report_path": path,
        },
        result={"report_path": path, "study_info": si.model_dump()},
    )


TOOLS = [
    Tool(
        "generate_272",
        "Generate an ICH M4E Module 2.7.2 (Summary of Clinical Pharmacology Studies) "
        "CTD-compliant DOCX from the current analysis state. Args: drug_name, sponsor, "
        "study_id, route, indication, pop_description, dose_range, matrix, assay_lloq "
        "(all optional strings; use whatever is known).",
        "regulatory",
        {
            "type": "object",
            "properties": {
                "drug_name":       {"type": "string"},
                "sponsor":         {"type": "string"},
                "study_id":        {"type": "string"},
                "route":           {"type": "string"},
                "indication":      {"type": "string"},
                "pop_description": {"type": "string"},
                "dose_range":      {"type": "string"},
                "matrix":          {"type": "string"},
                "assay_lloq":      {"type": "string"},
            },
            "required": [],
        },
        generate_272,
    ),
]
